from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import os

app = Flask(__name__)
CORS(app)

# json_data = {
#     "lista1": [
#         {"itens": "x", "DSCServico": "descriçao aaaaaa", "VCUnid": 1, "VCQuant": 4, "VCVU": 7, "VCValorTotal": 1},
#         {"itens": "y", "DSCServico": "descriçao aaaaaaaa", "VCUnid": 6, "VCQuant": 9, "VCVU": 6, "VCValorTotal": 34},
#         {"itens": "z", "DSCServico": "descriçao aaaaaaaaaaaa", "VCUnid": 3, "VCQuant": 6, "VCVU": 12, "VCValorTotal": 9},
#     ]
# }

def main():
   template_file_path = "Proposta.docx"
   output_file_path = 'Proposta NEW.docx'

   variables = {
       "${TXTRS}": "Nome razão social do ramon",
       "${NOMREPRESENTANTECLIENTE}": "P4pro "
    #    "${TABELA}": json_data
   }

   template_document = Document(template_file_path)

   for variable_key, variable_value in variables.items():
       for paragraph in template_document.paragraphs:
           replace_text_in_paragraph(paragraph, variable_key, variable_value)

    #    for table in template_document.tables:
    #        for col in table.columns:
    #            for cell in col.cells:
    #               for paragraph in cell.paragraphs:
    #                   replace_text_in_paragraph(paragraph, variable_key, variable_value)

   template_document.save(output_file_path)

def replace_text_in_paragraph(paragraph, key, value):
   if key in paragraph.text:
       inline = paragraph.runs
       for item in inline:
           if key in item.text:
               item.text = item.text.replace(key, value)

@app.route('/api/replace-word', methods=['POST'])
def replace_word_values():
   try:
       main()
       return jsonify({"status": "success", "message": "Sucesso, arquivo criado"})
   except Exception as error_mensage:
       return jsonify({"status": "error", "message": str(error_mensage)})

if __name__ == '__main__':
  app.run(debug=True)
